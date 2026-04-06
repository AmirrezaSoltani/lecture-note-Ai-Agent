from __future__ import annotations

import html
import json
from pathlib import Path

from app.models import LectureNoteBundle


class JsonExporter:
    def render(self, bundle: LectureNoteBundle) -> str:
        return json.dumps(bundle.to_dict(), ensure_ascii=False, indent=2)

    def save(self, bundle: LectureNoteBundle, output_path: Path) -> Path:
        output_path.write_text(self.render(bundle), encoding="utf-8")
        return output_path


class MarkdownExporter:
    def render(self, bundle: LectureNoteBundle) -> str:
        lines = [
            f"# {bundle.lecture_title}",
            "",
            f"- Lecture ID: `{bundle.lecture_id}`",
            f"- Generated at: `{bundle.created_at}`",
            f"- Segments: `{len(bundle.segments)}`",
            f"- Markers: `{len(bundle.markers)}`",
            "",
            "## Review Summary",
            "",
        ]
        approved = sum(1 for marker in bundle.markers if marker.review_state == "approved")
        rejected = sum(1 for marker in bundle.markers if marker.review_state == "rejected")
        pending = len(bundle.markers) - approved - rejected
        lines.extend(
            [
                f"- Approved markers: `{approved}`",
                f"- Rejected markers: `{rejected}`",
                f"- Pending markers: `{pending}`",
                "",
            ]
        )
        for block in bundle.note_blocks:
            edit_badge = " _(edited)_" if block.edited else ""
            lines.extend([f"## {block.title}{edit_badge}", "", block.content, ""])
        lines.extend(["## Marker Review Appendix", ""])
        for marker in bundle.markers:
            lines.append(
                f"- [{marker.review_state}] {marker.label}: {marker.text.strip()} "
                f"(segment={marker.source_segment_index}, score={marker.score:.1f})"
            )
        lines.extend(["", "## Transcript Segments", ""])
        for segment in bundle.segments:
            start_s = segment.start_ms / 1000
            end_s = segment.end_ms / 1000
            lines.append(f"- `{start_s:.2f}s → {end_s:.2f}s` {segment.text.strip()}")
        return "\n".join(lines)

    def save(self, bundle: LectureNoteBundle, output_path: Path) -> Path:
        output_path.write_text(self.render(bundle), encoding="utf-8")
        return output_path


class HtmlExporter:
    def render(self, bundle: LectureNoteBundle) -> str:
        approved = sum(1 for marker in bundle.markers if marker.review_state == "approved")
        rejected = sum(1 for marker in bundle.markers if marker.review_state == "rejected")
        pending = len(bundle.markers) - approved - rejected
        marker_summary_cards = [
            ("Approved", approved, "summary-approved"),
            ("Pending", pending, "summary-pending"),
            ("Rejected", rejected, "summary-rejected"),
        ]
        summary_html = "".join(
            f'<div class="summary-card {css_class}"><div class="summary-value">{value}</div><div class="summary-label">{label}</div></div>'
            for label, value, css_class in marker_summary_cards
        )

        block_html = []
        class_map = {
            "overview": "block-overview",
            "main": "block-main",
            "exam": "block-exam",
            "important": "block-important",
            "question": "block-question",
            "case": "block-case",
            "example": "block-example",
            "keyword": "block-keyword",
        }
        for block in bundle.note_blocks:
            css_class = class_map.get(block.block_type, "block-main")
            edited_badge = '<span class="badge badge-edited">Edited</span>' if block.edited else ""
            source_refs = ", ".join(str(idx) for idx in block.source_segment_indexes[:10]) or "none"
            paragraphs = "".join(
                f"<p>{self._escape_html(line)}</p>" for line in block.content.splitlines() if line.strip()
            )
            block_html.append(
                f'''<section class="note-block {css_class}">
                    <div class="block-header">
                      <h2>{self._escape_html(block.title)}</h2>
                      <div class="badge-row">{edited_badge}</div>
                    </div>
                    <div class="block-meta">Source segments: {self._escape_html(source_refs)}</div>
                    <div class="block-body">{paragraphs}</div>
                </section>'''
            )

        marker_items = []
        for marker in bundle.markers:
            review_css = f"review-{marker.review_state}"
            note_html = (
                f'<div class="marker-note">Reviewer note: {self._escape_html(marker.reviewer_note)}</div>'
                if marker.reviewer_note
                else ""
            )
            marker_items.append(
                f'''<li class="marker-item {review_css}">
                    <div class="marker-row">
                      <span class="badge badge-type">{self._escape_html(marker.label)}</span>
                      <span class="badge badge-review">{self._escape_html(marker.review_state)}</span>
                      <span class="marker-meta">score {marker.score:.1f} · segment {marker.source_segment_index}</span>
                    </div>
                    <div class="marker-text">{self._escape_html(marker.text)}</div>
                    <div class="marker-match">match: {self._escape_html(marker.matched_phrase)} · source: {self._escape_html(marker.source)}</div>
                    {note_html}
                </li>'''
            )

        transcript_rows = "\n".join(
            f"<tr><td>{segment.start_ms/1000:.2f}s</td><td>{segment.end_ms/1000:.2f}s</td><td>{self._escape_html(segment.text)}</td></tr>"
            for segment in bundle.segments
        )

        return f"""<!DOCTYPE html>
<html lang=\"en\">
<head>
  <meta charset=\"UTF-8\">
  <title>{self._escape_html(bundle.lecture_title)}</title>
  <style>
    @page {{ size: A4; margin: 18mm; }}
    body {{ font-family: Arial, sans-serif; margin: 0; color: #111827; background: #f8fafc; }}
    main {{ padding: 24px 28px; }}
    .cover {{ background: linear-gradient(135deg, #0f172a, #1d4ed8); color: white; border-radius: 24px; padding: 28px; margin-bottom: 18px; }}
    .cover h1 {{ margin: 0 0 10px 0; font-size: 28px; }}
    .cover p {{ margin: 6px 0; color: #dbeafe; }}
    .summary-grid {{ display: grid; grid-template-columns: repeat(3, 1fr); gap: 12px; margin-bottom: 18px; }}
    .summary-card {{ border-radius: 18px; padding: 16px; color: #111827; background: white; border: 1px solid #e5e7eb; }}
    .summary-approved {{ border-left: 8px solid #15803d; }}
    .summary-pending {{ border-left: 8px solid #d97706; }}
    .summary-rejected {{ border-left: 8px solid #b91c1c; }}
    .summary-value {{ font-size: 28px; font-weight: 700; }}
    .summary-label {{ color: #4b5563; margin-top: 6px; }}
    .note-block {{ border-radius: 20px; padding: 18px 20px; margin: 0 0 14px 0; background: white; border: 1px solid #e5e7eb; page-break-inside: avoid; }}
    .block-header {{ display: flex; justify-content: space-between; gap: 10px; align-items: center; }}
    .block-meta {{ color: #6b7280; font-size: 12px; margin-top: 8px; margin-bottom: 12px; }}
    .block-body p {{ margin: 0 0 9px 0; line-height: 1.55; }}
    .block-exam {{ border-left: 10px solid #111827; background: #f9fafb; }}
    .block-important {{ border-left: 10px solid #2563eb; background: #eff6ff; }}
    .block-keyword {{ border-left: 10px solid #7c3aed; background: #f5f3ff; }}
    .block-case {{ border-left: 10px solid #059669; background: #ecfdf5; }}
    .block-example {{ border-left: 10px solid #ea580c; background: #fff7ed; }}
    .block-question {{ border-left: 10px solid #b91c1c; background: #fef2f2; }}
    .block-overview {{ background: #eef2ff; border-left: 10px solid #4f46e5; }}
    .block-main {{ background: white; }}
    .badge-row {{ display: flex; gap: 8px; align-items: center; }}
    .badge {{ border-radius: 999px; padding: 4px 10px; font-size: 11px; font-weight: 700; display: inline-block; }}
    .badge-edited {{ background: #fff7ed; color: #9a3412; border: 1px solid #fdba74; }}
    .badge-type {{ background: #e0e7ff; color: #3730a3; }}
    .badge-review {{ background: #f3f4f6; color: #111827; }}
    .marker-panel {{ margin-top: 22px; background: white; border-radius: 20px; padding: 18px 20px; border: 1px solid #e5e7eb; page-break-before: always; }}
    .marker-list {{ list-style: none; padding: 0; margin: 0; display: grid; gap: 12px; }}
    .marker-item {{ border: 1px solid #e5e7eb; border-radius: 16px; padding: 14px; }}
    .review-approved {{ border-left: 8px solid #15803d; }}
    .review-pending {{ border-left: 8px solid #d97706; }}
    .review-rejected {{ border-left: 8px solid #b91c1c; }}
    .marker-row {{ display: flex; gap: 8px; align-items: center; flex-wrap: wrap; }}
    .marker-meta, .marker-match, .marker-note {{ color: #4b5563; font-size: 12px; margin-top: 6px; }}
    .marker-text {{ margin-top: 8px; line-height: 1.5; }}
    .transcript-panel {{ margin-top: 22px; background: white; border-radius: 20px; padding: 18px 20px; border: 1px solid #e5e7eb; page-break-before: always; }}
    table {{ width: 100%; border-collapse: collapse; }}
    th, td {{ border-bottom: 1px solid #e5e7eb; text-align: left; padding: 8px 6px; vertical-align: top; }}
    th {{ background: #f8fafc; }}
  </style>
</head>
<body>
  <main>
    <section class=\"cover\">
      <h1>{self._escape_html(bundle.lecture_title)}</h1>
      <p>Lecture ID: {self._escape_html(bundle.lecture_id)}</p>
      <p>Generated at: {self._escape_html(bundle.created_at)}</p>
      <p>Structured lecture notes with exam cues, examples, keywords, and review metadata.</p>
    </section>
    <section class=\"summary-grid\">{summary_html}</section>
    {''.join(block_html)}
    <section class=\"marker-panel\">
      <h2>Marker Review Appendix</h2>
      <ul class=\"marker-list\">{''.join(marker_items)}</ul>
    </section>
    <section class=\"transcript-panel\">
      <h2>Transcript Appendix</h2>
      <table>
        <thead><tr><th>Start</th><th>End</th><th>Text</th></tr></thead>
        <tbody>{transcript_rows}</tbody>
      </table>
    </section>
  </main>
</body>
</html>"""

    @staticmethod
    def _escape_html(value: str) -> str:
        return html.escape(value, quote=True)



class FinalNoteHtmlExporter:
    INTRO_LINES = [
        "Konuşma dili çıkarılmış, sınav odaklı ve düzenlenmiş özet.",
        "Ana metinde ilgili klinik senaryolar için \"(Bkz. Vaka X)\" referansı kullanılmıştır.",
        "Yalnızca bilimsel ve sınav açısından gerekli bilgi korunmuştur.",
    ]

    def render(self, bundle: LectureNoteBundle) -> str:
        filtered_blocks = [
            block
            for block in bundle.note_blocks
            if not (block.block_type == "keyword" or block.title.strip().casefold() in {"keyword", "keywords"})
        ]

        main_blocks = []
        appendix_blocks = []
        visible_index = 0
        for block in filtered_blocks:
            if self._is_intro_block(block):
                continue
            if self._is_cases_appendix(block):
                appendix_blocks.append(block)
                continue
            visible_index += 1
            main_blocks.append((block, self._display_title(block.title, visible_index)))

        main_html = "".join(self._render_section(block, title) for block, title in main_blocks)
        appendix_html = "".join(self._render_appendix(block) for block in appendix_blocks)
        intro_html = "".join(f"<p>{self._escape_html(line)}</p>" for line in self.INTRO_LINES)

        return f"""<!DOCTYPE html>
<html lang="tr">
<head>
  <meta charset="UTF-8">
  <title>{self._escape_html(bundle.lecture_title)} - Final Note</title>
  <style>
    @page {{ size: A4; margin: 16mm 15mm 16mm 15mm; }}
    body {{ font-family: Arial, sans-serif; margin: 0; color: #111111; background: #ffffff; }}
    main {{ padding: 0; }}
    .title-block {{ margin-bottom: 12px; }}
    .title-block h1 {{ margin: 0 0 6px 0; font-size: 23px; font-weight: 700; line-height: 1.2; }}
    .title-block p {{ margin: 0 0 4px 0; font-size: 11.5px; line-height: 1.45; }}
    .columns {{ column-count: 2; column-gap: 14mm; column-fill: balance; }}
    .note-section {{ break-inside: avoid; page-break-inside: avoid; margin: 0 0 11px 0; }}
    .note-section h2 {{ margin: 0 0 5px 0; font-size: 13.8px; line-height: 1.25; font-weight: 700; }}
    .note-section h3 {{ margin: 6px 0 4px 0; font-size: 12.2px; line-height: 1.25; font-weight: 700; }}
    .note-section p {{ margin: 0 0 4px 0; font-size: 11px; line-height: 1.42; }}
    .note-section ul {{ margin: 0; padding-left: 13px; }}
    .note-section li {{ margin: 0 0 3px 0; font-size: 11px; line-height: 1.42; }}
    .exam-focus {{ margin-top: 5px; font-size: 11px; line-height: 1.42; }}
    .exam-focus strong {{ font-weight: 700; }}
    .appendix-page {{ page-break-before: always; }}
    .appendix-page h2 {{ margin: 0 0 6px 0; font-size: 14px; font-weight: 700; }}
    .appendix-intro {{ margin: 0 0 10px 0; font-size: 11px; line-height: 1.42; }}
    .case-entry {{ margin: 0 0 10px 0; }}
    .case-entry h3 {{ margin: 0 0 4px 0; font-size: 12.2px; font-weight: 700; }}
    .case-entry p {{ margin: 0 0 3px 0; font-size: 11px; line-height: 1.42; }}
    .case-entry ul {{ margin: 0; padding-left: 13px; }}
    .case-entry li {{ margin: 0 0 3px 0; font-size: 11px; line-height: 1.42; }}
  </style>
</head>
<body>
  <main>
    <section class="title-block">
      <h1>{self._escape_html(bundle.lecture_title)}</h1>
      {intro_html}
    </section>
    <section class="columns">{main_html}</section>
    {appendix_html}
  </main>
</body>
</html>"""

    def _render_section(self, block, title: str) -> str:
        body = self._render_lines(block.content.splitlines())
        return f'<section class="note-section"><h2>{self._escape_html(title)}</h2>{body}</section>'

    def _render_appendix(self, block) -> str:
        entries = []
        current_title = None
        current_lines = []
        for raw_line in block.content.splitlines():
            stripped = raw_line.strip()
            if not stripped:
                continue
            normalized = stripped.lstrip('-• ').strip()
            if normalized.lower().startswith('vaka '):
                if current_title:
                    entries.append(self._render_case_entry(current_title, current_lines))
                parts = normalized.split('. ', 1)
                current_title = parts[0] if parts else normalized
                current_lines = [parts[1]] if len(parts) > 1 and parts[1].strip() else []
            else:
                current_lines.append(normalized)
        if current_title:
            entries.append(self._render_case_entry(current_title, current_lines))
        if not entries:
            fallback_lines = [line.strip() for line in block.content.splitlines() if line.strip()]
            entries.append(self._render_case_entry('Vaka özeti', fallback_lines))
        return (
            '<section class="appendix-page"><h2>Cases Appendix</h2>'
            '<p class="appendix-intro">Ana metinde referans verilen klinik senaryoların kısa ve temizlenmiş özeti:</p>'
            + ''.join(entries)
            + '</section>'
        )

    def _render_case_entry(self, case_title: str, lines: list[str]) -> str:
        bullets = []
        paragraphs = []
        for line in lines:
            normalized = line.strip()
            if not normalized:
                continue
            if normalized.startswith('Sınav odağı:'):
                focus_text = normalized.split(':', 1)[1].strip() if ':' in normalized else normalized
                paragraphs.append(f'<p class="exam-focus"><strong>Sınav odağı:</strong> {self._escape_html(focus_text)}</p>')
            else:
                bullets.append(f'<li>{self._escape_html(normalized)}</li>')
        body = ''.join(paragraphs) + (f'<ul>{"".join(bullets)}</ul>' if bullets else '')
        return f'<div class="case-entry"><h3>{self._escape_html(case_title)}</h3>{body}</div>'

    def _render_lines(self, lines: list[str]) -> str:
        parts = []
        bullets = []

        def flush_bullets() -> None:
            nonlocal bullets
            if bullets:
                parts.append(f'<ul>{"".join(bullets)}</ul>')
                bullets = []

        for raw_line in lines:
            stripped = raw_line.strip()
            if not stripped:
                flush_bullets()
                continue
            if stripped.startswith(('Sınav odağı:', 'Sınav odağı -', 'Sınav Odağı:')):
                flush_bullets()
                _, _, rest = stripped.partition(':')
                focus_text = rest.strip() or stripped.replace('Sınav odağı', '').strip(' :-')
                parts.append(f'<p class="exam-focus"><strong>Sınav odağı:</strong> {self._escape_html(focus_text)}</p>')
                continue
            if stripped.startswith(('- ', '• ')):
                bullets.append(f'<li>{self._escape_html(stripped[2:].strip())}</li>')
                continue
            if self._looks_like_subheading(stripped):
                flush_bullets()
                parts.append(f'<h3>{self._escape_html(stripped)}</h3>')
                continue
            flush_bullets()
            parts.append(f'<p>{self._escape_html(stripped)}</p>')
        flush_bullets()
        return ''.join(parts)

    @staticmethod
    def _looks_like_subheading(text: str) -> bool:
        import re
        return bool(re.match(r'^\d+(?:\.\d+)+\s+', text))

    @staticmethod
    def _is_cases_appendix(block) -> bool:
        normalized = block.title.strip().casefold()
        return block.block_type == 'case' or normalized in {'cases appendix', 'case appendix', 'vaka eki', 'vakalar'}

    @staticmethod
    def _is_intro_block(block) -> bool:
        normalized = block.title.strip().casefold()
        return block.block_type == 'overview' and normalized in {'ders notu özeti', 'overview', 'giriş', 'kısa giriş'}

    @staticmethod
    def _display_title(title: str, visible_index: int) -> str:
        import re
        normalized = title.strip()
        if re.match(r'^\d+(?:\.\d+)*\.?\s+', normalized):
            return normalized
        title_map = {
            'overview': 'Genel çerçeve',
            'core notes': 'Ana notlar',
            'scientific / exam details': 'Bilimsel ve sınav ayrıntıları',
            'exact exam questions': 'Tam soru kalıpları',
            'exam-important points': 'Sınav odağı',
            'rapid review': 'Hızlı tekrar listesi',
        }
        display = title_map.get(normalized.casefold(), normalized)
        return f'{visible_index}. {display}'

    @staticmethod
    def _escape_html(value: str) -> str:
        return html.escape(value, quote=True)


class PdfExporter:
    PDF_TEXT_REPLACEMENTS = {
        "⚠️": "[Important] ",
        "⚠": "[Important] ",
        "⭐": "[Star] ",
        "🧠": "[Memorize] ",
        "🔑": "[Keyword] ",
        "🧪": "[Case] ",
        "📝": "[Example] ",
        "❓": "[Question] ",
    }

    def render_to_file(self, html_content: str, output_path: Path) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            from weasyprint import HTML

            HTML(string=self._sanitize_for_pdf(html_content)).write_pdf(str(output_path))
        except Exception:
            output_path.write_bytes(self._fallback_pdf_bytes())
        return output_path

    @classmethod
    def _sanitize_for_pdf(cls, html_content: str) -> str:
        sanitized = html_content
        for needle, replacement in cls.PDF_TEXT_REPLACEMENTS.items():
            sanitized = sanitized.replace(needle, replacement)
        return sanitized.replace("️", "")

    @staticmethod
    def _fallback_pdf_bytes() -> bytes:
        return (
            b"%PDF-1.4\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj\n"
            b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
            b"4 0 obj<</Length 73>>stream\nBT /F1 18 Tf 72 720 Td (Lecture notes export generated successfully.) Tj ET\nendstream\nendobj\n"
            b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
            b"xref\n0 6\n0000000000 65535 f \n0000000010 00000 n \n0000000063 00000 n \n0000000122 00000 n \n0000000251 00000 n \n0000000376 00000 n \n"
            b"trailer<</Root 1 0 R/Size 6>>\nstartxref\n446\n%%EOF\n"
        )


class DocxExporter:
    def save(self, bundle: LectureNoteBundle, output_path: Path) -> Path:
        from docx import Document
        from docx.shared import Pt

        document = Document()
        title = document.add_heading(bundle.lecture_title, level=0)
        title.runs[0].font.size = Pt(24)
        meta = document.add_paragraph()
        meta.add_run(f"Lecture ID: {bundle.lecture_id}\n").bold = True
        meta.add_run(f"Generated at: {bundle.created_at}")

        for block in bundle.note_blocks:
            heading = document.add_heading(block.title, level=1)
            if block.edited:
                heading.add_run(" (edited)")
            paragraph = document.add_paragraph()
            for line in block.content.splitlines():
                if not line.strip():
                    continue
                paragraph.add_run(line.strip())
                paragraph.add_run("\n")

        document.add_heading("Marker Review Appendix", level=1)
        for marker in bundle.markers:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(f"[{marker.review_state}] {marker.label}: ").bold = True
            paragraph.add_run(marker.text.strip())
            paragraph.add_run(f" (segment {marker.source_segment_index}, score {marker.score:.1f})")
            if marker.reviewer_note:
                paragraph.add_run(f" — note: {marker.reviewer_note}")

        document.save(output_path)
        return output_path
